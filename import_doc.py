from docx import Document, styles
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import os
from glob_var import doc_TableHeader, doc_Titles, doc_FirstCol, doc_SecCol, doc_ThrdCol, DESC_DIM, doc_MonthNames
import re
import time


def create_DocumentTitle(document):
    docTitle = document.add_paragraph()
    docTitle_run = docTitle.add_run('Технический дизайн\n')
    docTitle_run.font.size = Pt(26)
    docTitle_run.bold = True

    docTitle_run = docTitle.add_run('на внедрение Автоматизированной Системы Планирования\n')
    docTitle_run.font.size = Pt(14)
    docTitle_run.bold = True

    docTitle_run = docTitle.add_run('«Название Компании»')
    docTitle_run.font.size = Pt(14)
    docTitle_run.bold = True

    docTitle_frm = docTitle.paragraph_format
    docTitle_frm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    docTitle_frm.space_before = Pt(150)
    docTitle_frm.space_after = Pt(200)

    cur_date = time.localtime()
    day, month, year = cur_date.tm_mday, cur_date.tm_mon, cur_date.tm_year
    month = doc_MonthNames[month]
    docDate = document.add_paragraph()
    docDate_run = docDate.add_run(f'«{day}» {month} {year} года')
    docDate_run.underline = True

    docDate_frm = docDate.paragraph_format
    docDate_frm.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    docDate_frm.space_after = Pt(100)

    docTail = document.add_paragraph()
    docTail_run = docTail.add_run('Москва')
    docTail_run.font.size = Pt(14)
    docTail_run.bold = True

    docTail_frm = docTail.paragraph_format
    docTail_frm.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()


def get_strdate():
    cur_datetime = time.localtime()
    pass


def create_TableContent(document):
    content = document.add_paragraph()
    content_run = content.add_run('Содержание:')
    content_run.font.size = Pt(20)
    # content_frm = content.paragraph_format
    # content_frm.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-7" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:updateFields')
    # fldChar3.set(qn('w:val'), 'true')
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    # p_element = paragraph._p

    document.add_page_break()


def set_bg_color(cell):
    """
    set background shading for Header Rows
    """
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), "D3D3D3")
    tblCellProperties.append(clShading)
    return cell


def import_Description(path):
    dict_desc = {}

    if os.path.exists(DESC_DIM):
        with open(path, 'r', encoding='cp1251') as f:
            for line in f:
                key, value = re.split(',', line.replace('"', ''), maxsplit=1)
                dict_desc[key] = value
    return dict_desc


def create_AppStructure(dict_app, document, lvl=1, num_pref='9.', start_number=1):
    tbl_header = doc_TableHeader['view']
    cols_info = (doc_FirstCol['view'], doc_SecCol['view'])
    titles = doc_Titles
    cols_cnt = len(tbl_header)
    doc_type = 'view'

    for folder, content in dict_app.items():
        # num_pref = f'{num_pref}{start_number}.'
        document.add_heading(f'{num_pref}{start_number}. {folder}', level=lvl)
        subfolders = content['subfolder']

        if len(subfolders) > 0:
            create_AppStructure(subfolders, document, lvl=(lvl+1), num_pref=f'{num_pref}{start_number}.')

        for n_, app in enumerate(content['app']):
            app_name, app_type, app_info = app.values()
            app_title = f'\u2022 {doc_Titles["app_name"]} "{app_name}"'
            app = document.add_heading(app_title, level=(lvl+1))
            app_format = app.paragraph_format
            app_format.left_indent = Cm(0.5)

            add_Table(document, 'view', app_info, cols_info)

        start_number += 1


def create_ProcessInfo(records, document, lvl=2, num_pref='10.', start_number=1):
    doc_type = 'process'
    cols_info = (doc_FirstCol[doc_type], doc_SecCol[doc_type])
    titles = doc_Titles['pro_name']

    process_use = []
    for n, record in enumerate(records):
        pro_name = record['pro_name']
        title = f'{n+1}. {titles} "{pro_name}"'
        pro = document.add_heading(title, level=(lvl+1))
        create_hyperlink(pro_name, title)
        pro_format = pro.paragraph_format
        pro_format.left_indent = Cm(0.5)

        add_Table(document, doc_type, record, cols_info)


def add_Table(document, doc_type, data, cols_info):
    merge_row = []
    cnt_row = 0
    frst_col, sec_col = cols_info

    table = create_Table(document, header=doc_TableHeader[doc_type])
    for row_key, row_name in frst_col.items():
        frst_txt, sec_txt, thr_txt = row_name, '-', '-'
        values = data[row_key]

        # считаем строки для последующего объединения
        # cnt_row = 0
        if isinstance(values, dict):
            for key_, value_ in values.items():
                sec_txt = sec_col[key_]
                if len(value_) == 0:
                    thr_txt = '-'
                elif isinstance(value_, list):
                    thr_txt = str.join('\n', value_)
                elif isinstance(value_, set):
                    thr_txt = str.join('\n', value_)
                else:
                    thr_txt = value_

                record = (frst_txt, sec_txt, thr_txt)
                add_row(table, record, type=doc_type)
                cnt_row += 1

        elif isinstance(values, list):
            if len(values) == 0:
                record = (frst_txt, sec_txt, thr_txt)
                add_row(table, record, type=doc_type)
                cnt_row += 1
            else:
                for num, value in enumerate(values):
                    if isinstance(value, str):
                        sec_txt = value
                    elif isinstance(value, list) or isinstance(value, tuple):
                        sec_txt, *dim_desc = value
                        if len(dim_desc) == 0:
                            thr_txt = dim_desc
                        if len(dim_desc) == 1:
                            thr_txt = dim_desc[0]
                        else:
                            thr_txt = str.join('\n', dim_desc)

                    record = (frst_txt, sec_txt, thr_txt)
                    add_row(table, record, type=doc_type)
                    cnt_row += 1

        else:
            if values == '':
                if doc_type == 'view':
                    continue
                else:
                    values = '-'
            if doc_type == 'view':
                thr_txt = doc_ThrdCol[doc_type][row_key]
            record = (frst_txt, values, thr_txt)
            add_row(table, record, type=doc_type)
            cnt_row += 1

        merge_row.append(cnt_row)

    table = format_row(table, merge_row)

    for i, col in enumerate(table.columns):
        if i == 0:
            col_width = Cm(5)
        else:
            col_width = Cm(6)
        col.width = col_width


def create_Table(document, header, style='Table Grid'):
    cols_cnt = len(header)

    table = document.add_table(rows=1, cols=cols_cnt)
    table.style = style
    hdr_cells = table.rows[0].cells

    for num, hdr_name in enumerate(header):
        hdr_cells[num].text = str(hdr_name)
        set_bg_color(hdr_cells[num])

    return table


def format_row(table, merge_row):
    for n, other_row in enumerate(merge_row[:-1]):
        start_ = other_row + 1
        stop_ = merge_row[n + 1]
        try:
            mrg_txt = table.cell(start_, 0).text
        except IndexError:
            continue
        if start_ != stop_:
            table.cell(start_, 0).merge(table.cell(stop_, 0))
            table.cell(start_, 0).text = mrg_txt
    return table


def add_row(table, record, type='view'):
    header, name, description = record
    if isinstance(name, set):
        name = ''.join(name)

    row_cells = table.add_row().cells
    row_cells[0].text = header
    set_bg_color(row_cells[0])

    if type == 'view':
        if header == 'Процессы' and len(name) > 0:
            p = row_cells[1].add_paragraph()
            p.add_run(name)
            try:
                process_links[name].append(p)
            except KeyError:
                process_links[name] = [p]
        else:
            row_cells[1].text = name
        row_cells[2].text = description
    elif type == 'process':
        # Костыль с измерением
        # Продумать на будущее формат для пустых значений
        if header == 'Приёмник:':
            row_cells[1].text = f'{name}: {description}'
        else:
            row_cells[1].text = name


def create_doc(records, doc_path, par_stl='List Bullet'):
    global process_links
    process_links = {}

    doc = Document()

    create_DocumentTitle(doc)
    create_TableContent(doc)

    num_pref = 1
    doc.add_heading(f'{num_pref}. Описание структуры системы', level=1)
    create_AppStructure(records[0], doc, lvl=2, num_pref=f'{num_pref}.')

    doc.add_page_break()

    num_pref = 2
    doc.add_heading(f'{num_pref}. Описанние процессов модели', level=1)
    create_ProcessInfo(records[1], doc, lvl=2, num_pref=f'{num_pref}.')

    doc.save(doc_path)

def create_hyperlink(name, title,):
    name = name.replace('}Drill_', '')
    try:
        for paragraph in process_links[name]:
            add_hyperlink(paragraph, title, name)
    except KeyError:
        pass

def add_hyperlink(paragraph, heading, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: A Run object containing the hyperlink
    """

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), heading, )

    # Create a w:r element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    # Join all the xml elements together add add the required text to the w:r element
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.clear().add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return r


if __name__ == '__main__':

    script_path = os.path.dirname(__file__)
    doc_name = 'demo.docx'
    doc_path = os.path.join(script_path, doc_name)
    doc_type = 'view'
    par_stl = 'List Bullet'


    # records = (records1, records2)
    # create_doc(records, doc_path, par_stl)




